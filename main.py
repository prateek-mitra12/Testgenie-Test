from config import config
import streamlit as st
from typing import List, Tuple, Dict
from role_prompt import role_prompt
import random
import json
import boto3
import pandas as pd
import PyPDF2
from pypdf import PdfReader
import docx
import os
import time
import uuid
import datetime
from chatHistory import *
from io import StringIO
from langchain_community.chat_message_histories import StreamlitChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from models import *
from rag_system import setup_rag_system
from streamlit_feedback import streamlit_feedback
from routeLLM import getRoutingModel
from routeLLM import model_id_name_mapping
from modelOfModels import modelOfModels


st.set_page_config(page_title=r"TestGenie",
                   layout="wide",
                   )
st.header(r" TestGenie 🤖")


def update_toggle(source):
    if source == "mom":
        st.session_state["mom_enabled"] = not st.session_state["mom_enabled"]
        st.session_state["routeLLM_enabled"] = False  # Disable RouteLLM
    elif source == "routeLLM":
        st.session_state["routeLLM_enabled"] = not st.session_state["routeLLM_enabled"]
        st.session_state["mom_enabled"] = False  # Disable MOM


def generate_session_id():
    """Generate a unique session ID"""
    return str(uuid.uuid4())


def generate_chat_title(first_message, max_length=50):
    """Generate a meaningful title from the first user message"""
    if not first_message:
        return "New Chat"
    
    # Remove "Human: " prefix if present
    if first_message.startswith("Human: "):
        first_message = first_message[7:]
    
    # Truncate and clean the message
    title = first_message.strip()[:max_length]
    if len(first_message) > max_length:
        title += "..."
    
    return title


def save_current_chat_history():
    """Save the current chat history to DynamoDB"""
    try:
        # Get the current tab's chat history
        current_tab = st.session_state.get("current_tab", "JSON Payload Test Cases")
        
        if current_tab == "JSON Payload Test Cases":
            history = st.session_state.get("chat_history", [])
        elif current_tab == "Test Cases (Gherkin Format)":
            history = st.session_state.get("chat_history_gherkin", [])
        elif current_tab == "Playwright Scripts":
            history = st.session_state.get("chat_history_playwright", [])
        elif current_tab == "Code Generator":
            history = st.session_state.get("chat_history_code_generator", [])
        elif current_tab == "Unit Testing":
            history = st.session_state.get("chat_history_unit_testing", [])
        else:
            history = []
        
        # Only save if there's actually some conversation
        if len(history) >= 2:  # At least one human message and one AI response
            session_id = st.session_state.get("session_id")
            if not session_id:
                session_id = generate_session_id()
                st.session_state["session_id"] = session_id
            
            # Generate title from first human message if not already set
            title = st.session_state.get("chat_history_title")
            if not title:
                first_human_message = next((msg for msg in history if msg.startswith("Human:")), "")
                title = generate_chat_title(first_human_message)
                st.session_state["chat_history_title"] = title
            
            # Save to DynamoDB
            save_chat_history(session_id, title, history)
            st.session_state["chat_history_saved"] = True
            print(f"Chat history saved successfully for session: {session_id}")
            
    except Exception as e:
        print(f"Error saving chat history: {str(e)}")
        st.error(f"Failed to save chat history: {str(e)}")


def render_sidebar() -> Tuple[Dict, int, str]:
    ####################################### Sidebar Navigation Controls #####################################################
    
    # New Chat Button (at the top)
    st.sidebar.button("New Chat", on_click=new_chat, type="primary", use_container_width=True)

    with st.sidebar:
        if "mom_enabled" not in st.session_state:
            st.session_state["mom_enabled"] = False
        if "routeLLM_enabled" not in st.session_state:
            st.session_state["routeLLM_enabled"] = False
        
        # Model Selection
        st.markdown("**Model**")
        model_name_select = st.selectbox(
            '',
            list(config["models"].keys()),
            key=f"{st.session_state['widget_key']}_Model_Id",
            disabled=st.session_state["routeLLM_enabled"] or st.session_state["mom_enabled"]
        )

        st.session_state["model_name"] = model_name_select

        model_config = config["models"][model_name_select]
        
    with st.sidebar.expander("Configure Model"): 
        role_select = st.selectbox(
            'Role',
            list(role_prompt.keys()) + ["Custom"],
            key=f"{st.session_state['widget_key']}_role_Id"
        )

        # Set the initial value of the text area based on the selected role
        role_prompt_text = "" if role_select == "Custom" else role_prompt.get(role_select, "")

        system_prompt_disabled = model_config.get("system_prompt_disabled", False)
        
        system_prompt = st.text_area(
            "System Prompt",
            value = role_prompt_text,
            key=f"{st.session_state['widget_key']}_System_Prompt",
            disabled=system_prompt_disabled,
        )

        with st.container():
            col1, col2 = st.columns(2)
            with col1:   
                web_local = st.selectbox(
                    'Options',
                    ('RAG', 'Web', 'Local'),
                    key=f"{st.session_state['widget_key']}_Options",
                )     
            with col2:  
                temperature = st.slider(
                    "Temperature",
                    min_value=0.0,
                    max_value=1.0,
                    value = model_config.get("temperature", 1.0),
                    key=f"{st.session_state['widget_key']}_Temperature",
                )
        with st.container():
            col1, col2 = st.columns(2)
            with col1:
                top_p = st.slider(
                    "Top-P",
                    min_value=0.0,
                    max_value=1.0,
                    value=model_config.get("top_p", 1.0),
                    step=0.01,
                    key=f"{st.session_state['widget_key']}_Top_P",
                )
            with col2:
                if "meta.llama2" in model_config["model_id"]:
                    max_tokens = st.slider(
                        "Max Token",
                        min_value=0,
                        max_value=2048,
                        value=model_config.get("max_tokens", 2048),  
                        step=8,
                        key=f"{st.session_state['widget_key']}_Max_Token",
                    )
                elif "amazon.titan-text-premier" in model_config["model_id"]:
                    max_tokens = st.slider(
                        "Max Token",
                        min_value=0,
                        max_value=3072,
                        value=model_config.get("max_tokens", 3072),  
                        step=8,
                        key=f"{st.session_state['widget_key']}_Max_Token",
                    )
                else:
                    max_tokens = st.slider(
                        "Max Token",
                        min_value=0,
                        max_value=4096,
                        value=model_config.get("max_tokens", 4096),  
                        step=8,
                        key=f"{st.session_state['widget_key']}_Max_Token",
                    )
        with st.container():
            if not "meta.llama" in model_config["model_id"] and not "amazon.titan" in model_config["model_id"]:
                col1, col2 = st.columns(2)
                with col1:
                    top_k = st.slider(
                        "Top-K",
                        min_value=1,
                        max_value=500,
                        value=model_config.get("top_k", 500),
                        step=5,
                        key=f"{st.session_state['widget_key']}_Top_K",
                    )
                with col2:
                    memory_window = st.slider(
                        "Memory Window",
                        min_value=0,
                        max_value=10,
                        value=model_config.get("memory_window", 10),
                        step=1,
                        key=f"{st.session_state['widget_key']}_Memory_Window",
                    )

    # Choose a tab Section
    st.sidebar.markdown("**Choose a tab**")
    st.session_state["current_tab"] = st.sidebar.selectbox("", ["JSON Payload Test Cases", "Test Cases (Gherkin Format)", "Playwright Scripts", "Code Generator", "Unit Testing"])

    st.sidebar.divider()  # Separator between controls and chat history

    ####################################### Display Chat History #####################################################

    histories = get_chat_histories()

    # CSS for equal button heights and truncating text
    st.markdown(
        """
        <style>
            .custom-button {
                display: flex;
                align-items: center;
                justify-content: center;
                height: 40px;
                overflow: hidden;
                white-space: nowrap;
                text-overflow: ellipsis;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Optional: Add Chat History section header
    st.sidebar.markdown("**Chat History**")

    @st.dialog("Chat History Options")
    def chatHistoryOptions(hist):
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("Delete Chat", key=f"delete_{hist['session_id']}", use_container_width=True):
                delete_chat_history(hist['session_id'])
                if st.session_state["session_id"] == hist['session_id']:
                    st.session_state.clear()
                st.rerun()
        with col2:
            if st.download_button(
                label="Download Chat History",
                data=download_chat_history_markdown(hist['session_id']),
                file_name=f"{hist['title'].replace(' ', '_')}.md",
                mime="text/markdown",
                key=f"download_{hist['session_id']}",
                use_container_width=True
            ):
                pass

        send_button_pressed = False
        
        col3, col4 = st.columns([5, 1])
        with col3:
            email_input = st.text_input("Share To (comma-separated emails)", key=f"email_{hist['session_id']}", placeholder="email1@example.com, email2@example.com")
        with col4:
            st.markdown("<div style='height: 28px;'></div>", unsafe_allow_html=True)  # Adjusts vertical alignment
            if st.button("Send", key=f"share_{hist['session_id']}"):
                send_button_pressed = True
                email_list = [email.strip() for email in email_input.split(",") if email.strip()]
                emails = email_list
                valid_emails = [email for email in emails if is_valid_email(email)]
                invalid_emails = [email for email in emails if not is_valid_email(email)]
                
        if send_button_pressed:
            if invalid_emails:
                st.warning(f"Invalid email addresses: {', '.join(invalid_emails)}")
            
            if valid_emails:
                with st.spinner("Sharing Chat History"):
                    success, error = share_chat_history(hist['session_id'], valid_emails)
                    if success:
                        st.success(f"✅ Chat History Shared Successfully !!")
                    else:
                        st.error(f"Error sending email: {error}")
            else:
                st.error("No valid email addresses provided !!")

            send_button_pressed = False
        
        new_title = st.text_input("✏️ Rename", hist['title'], key=f"rename_{hist['session_id']}")
        if new_title and new_title != hist['title']:
            rename_chat_history(hist['session_id'], new_title)
            st.success("Title Rename Done")
            st.rerun()
        last_updated_at = f"Last Updated : {datetime.datetime.fromtimestamp(int(hist['last_updated'])).strftime('%Y-%m-%d %H:%M:%S')}"
        st.markdown(f"<div style='text-align: right;'>{last_updated_at}</div>", unsafe_allow_html=True)

    # Display available histories in sidebar
    for hist in histories:
        # Truncate title for button display
        max_title_length = 20  # Adjust as needed
        truncated_title = hist['title'] if len(hist['title']) <= max_title_length else hist['title'][:max_title_length] + "..."
        button_label = f"{truncated_title}"
        
        col1, col2 = st.sidebar.columns([3, 1])
        with col1:
            if st.button(button_label, key=f"select_{hist['session_id']}", use_container_width=True):
                selected_history = get_chat_history(hist['session_id'])
                st.session_state["chat_history"] = selected_history
                st.session_state["session_id"] = hist['session_id']
                st.session_state["chat_history_title"] = hist['title']
                st.session_state["chat_history_saved"] = True
        with col2:
            if st.button("...", key=f"select_{hist['session_id']}_chat_history_options", use_container_width=True):
                chatHistoryOptions(hist)

    #################################################################################################################

    top_k = locals().get('top_k', 0)
    memory_window = locals().get('memory_window', 0)

    if "meta.llama2" in model_config["model_id"]:
        model_kwargs = {
            "temperature": temperature,
            "top_p": top_p,
            "max_gen_len": max_tokens,
        }
    elif "amazon.titan-text-premier" in model_config["model_id"]:
        model_kwargs = {
            "temperature": temperature,
            "topP": top_p,
            "maxTokenCount": max_tokens,
        }
    elif "anthropic.claude" in model_config["model_id"]:
        model_kwargs = {
            "max_tokens": 200000,
        }
    else:
        model_kwargs = {
            "temperature": temperature,
            "top_p": top_p,
            "top_k": top_k if top_k is not None else 0,
            "max_tokens": max_tokens,
        }

    if model_config.get("system_prompt_disabled", False):
        model_kwargs["system"] = system_prompt

    if "meta.llama2" in model_config["model_id"]:
        return model_kwargs, web_local
    
    return model_kwargs, memory_window, web_local


def get_session_history_streamlit_chat_messsages() -> StreamlitChatMessageHistory:
    history = StreamlitChatMessageHistory(key = 'messages')
    return history


def init_conversationchain(chat_model: ChatModel) -> RunnableWithMessageHistory:
    """
    Initialize the ConversationChain with the given parameters.
    """
    chain = RunnableWithMessageHistory(chat_model.llm, get_session_history_streamlit_chat_messsages, input_messages_key="question", history_messages_key="history",)
    return chain


def generate_response(
    chain: RunnableWithMessageHistory, input: str, chatmodel: ChatModel, document_uploaded: bool, tab: int) -> str:
    """
    Generate a response from the conversation chain with the given input.
    """
    sys = ""
    if document_uploaded:
        sys = chatmodel.format_prompt(tab)
    config = {"configurable": {"session_id": "any"}}

    return chain.stream({"question":input,"system":sys}, config)


def generate_response_titan(
    chain: RunnableWithMessageHistory, input: str, chatmodel: ChatModel, document_uploaded: bool) -> str:
    """
    Generate a response from the conversation chain with the given input.
    """
    sys = ""
    if document_uploaded:
        sys = chatmodel.format_prompt()
    config = {"configurable": {"session_id": "any"}}

    return chain.invoke({"question":input,"system":sys}, config)


def use_ai21(model_kwargs, prompt, chat_model, document_uploaded):
    # Initialize the Bedrock client
    client = boto3.client('bedrock-runtime', region_name='us-east-1')

    # Define your request parameters
    model_id = "ai21.j2-ultra-v1"
    sys = ""
    if document_uploaded:
        sys = chat_model.format_prompt()
    prompt_updated = json.dumps({"question":prompt,"system":sys})

    # Create the request payload
    payload = {
        "prompt": prompt_updated,
        "maxTokens": model_kwargs["max_tokens"],
        "temperature": model_kwargs["temperature"],
        "topP": model_kwargs["top_p"],
        "stopSequences": [],
        "countPenalty": {"scale": 0},
        "presencePenalty": {"scale": 0},
        "frequencyPenalty": {"scale": 0}
    }

    # Make the API request
    response = client.invoke_model(
        modelId=model_id,
        body=json.dumps(payload),
        contentType="application/json",
        accept="application/json"
    )

    # Process the response
    result = json.loads(response['body'].read())
    response_text = result["completions"][0]["data"]["text"]
    return response_text


def new_chat() -> None:
    """
    Reset the chat session and initialize a new conversation chain.
    """
    # Save current chat before starting new one
    save_current_chat_history()
    
    # Clear all session state related to chat
    st.session_state["chat_history"] = []
    st.session_state["chat_history_gherkin"] = []
    st.session_state["chat_history_playwright"] = []
    st.session_state["chat_history_code_generator"] = []
    st.session_state["chat_history_unit_testing"] = []
    st.session_state["rag_chain"] = None
    st.session_state["file_uploader_key"] = 0
    
    # Reset session ID and title for new chat
    st.session_state["session_id"] = None
    st.session_state["chat_history_title"] = None
    st.session_state["chat_history_saved"] = False


def excel_sheets_to_csv(uploaded_file):
    # Load the Excel file
    xls = pd.ExcelFile(uploaded_file)
    
    # Dictionary to store CSV content of each sheet
    sheet_csv_dict = {}

    # Loop through each sheet and convert it to CSV format
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        
        # Convert DataFrame to CSV
        csv_buffer = StringIO()
        df.to_csv(csv_buffer, index=False)
        
        # Store the CSV content for each sheet
        sheet_csv_dict[sheet_name] = csv_buffer.getvalue()

    # Combine the CSVs for all sheets in a readable format
    document_content = ""
    for sheet_name, csv_content in sheet_csv_dict.items():
        document_content += f"Sheet: {sheet_name}\n"
        document_content += csv_content
        document_content += "\n\n"

    return document_content


def fbcb(response):
    last_entry = st.session_state.chat_history[-1]  # get the last entry

    # {"type":"thumbs","score":"👍","text":"good bot"}
    feedback = f"\n\nFeedback: {{Score: {response['score']}, Text: {response['text']}}}"
    last_entry = f"{last_entry} {feedback}"
    st.session_state.chat_history[-1] = last_entry  

    st.write(feedback)

    # Create a new feedback by changing the key of feedback component.
    st.session_state.fbk = str(uuid.uuid4())


def stream_response(response_text):
    """Convert response text into a generator for streaming."""
    for word in response_text.split():
        yield word + " "  # Yield words one by one with a space


# Add placeholder functions for missing imports
def getSystemPrompt():
    return "You are a helpful assistant for generating JSON payload test cases."

def getSystemPromptGherkinFormat():
    return "You are a helpful assistant for generating test cases in Gherkin format."

def getSystemPromptPlaywright():
    return "You are a helpful assistant for generating Playwright scripts."

def getSystemPromptCodeGenerator():
    return "You are a helpful assistant for generating code."

def getSystemPromptUnitTesting():
    return "You are a helpful assistant for generating unit tests."


def handle_chat_interaction(prompt, tab_name, history_key, system_prompt_func, tab_number=None):
    """
    Centralized function to handle chat interactions across all tabs
    """
    # Ensure session ID exists
    if "session_id" not in st.session_state or not st.session_state["session_id"]:
        st.session_state["session_id"] = generate_session_id()
    
    # Get current model configuration
    model_kwargs = st.session_state.get("model_kwargs", {})
    
    # Get RAG prompt if available
    rag_prompt = ""
    if "rag_prompt" in st.session_state:
        rag_prompt = st.session_state["rag_prompt"]
    
    # Initialize chat model and conversation chain
    chat_model = ChatModel(st.session_state["model_name"], model_kwargs, rag_prompt)
    conv_chain = init_conversationchain(chat_model)
    
    # Determine which model to route to
    modelRoutedTo = st.session_state["model_name"]
    if st.session_state.get("routeLLM_enabled", False):
        if rag_prompt:
            model_ID_routedTo = getRoutingModel(f"User Query : {prompt} \n\n Uploaded Document : {rag_prompt} \n\n {system_prompt_func()}")
        else:
            model_ID_routedTo = getRoutingModel(prompt)
        modelRoutedTo = model_id_name_mapping[model_ID_routedTo]
        chat_model = ChatModel(modelRoutedTo, model_kwargs, rag_prompt)
        conv_chain = init_conversationchain(chat_model)

    st.session_state.question_state = True

    # Add user message to chat
    st.chat_message("human").write(prompt)
    st.session_state[history_key].append(f"Human: {prompt}")

    # Generate and display AI response
    if modelRoutedTo == "AI21 J2-Ultra":
        if rag_prompt:
            response = use_ai21(model_kwargs, prompt, chat_model, True)
        else:
            response = use_ai21(model_kwargs, prompt, chat_model, False)
        st.chat_message("ai").write(response)
        st.session_state[history_key].append(f"AI: {response}")
    elif modelRoutedTo == "amazon.titan":
        if rag_prompt:
            response = generate_response_titan(conv_chain, prompt, chat_model, True)
        else:
            response = generate_response_titan(conv_chain, prompt, chat_model, False)
        st.chat_message("ai").write(response.content)
        st.session_state[history_key].append(f"AI: {response.content}")
    else:
        with st.chat_message("ai"):
            if st.session_state.get("mom_enabled", False):
                if rag_prompt:
                    response = modelOfModels(f"User Query : {prompt} \n\n Uploaded Document : {rag_prompt} \n\n {system_prompt_func()}")
                else:
                    response = modelOfModels(prompt)
                
                st.write_stream(stream_response(response))
                st.session_state[history_key].append(f"AI: {response}")
            else:
                if rag_prompt and tab_number:
                    response = st.write_stream(generate_response(conv_chain, prompt, chat_model, True, tab_number))
                else:
                    response = st.write_stream(generate_response(conv_chain, prompt, chat_model, False, tab_number or 1))
                
                 #############Auto looping logic#############################################
                check = """
                    \n\n
                    "Please end your response with one of these indicators on a new line:\n"
                    "[RESPONSE_COMPLETE] - if your answer is fully complete\n"
                    "[RESPONSE_INCOMPLETE] - if you need to continue"
                """
                if rag_prompt:
                    response = st.write_stream(generate_response(conv_chain, prompt + check, chat_model, True, tab_number))
                else:
                    response = st.write_stream(generate_response(conv_chain, prompt + check, chat_model, False, tab_number))
                # Get continuations if needed
                while True:
                    if "[RESPONSE_COMPLETE]" in response:
                        # Clean up and we're done
                        response = response.replace("[RESPONSE_COMPLETE]", "").strip()
                        break
                    else:
                        # if "[RESPONSE_INCOMPLETE]" in response or If no indicator found, assume incomplete
                        # Clean up current portion
                        response = response.replace("[RESPONSE_INCOMPLETE]", "").strip()
                       
                        # Get continuation
                        continuation_prompt = (
                            "Please continue your previous response without repeating any information.\n\n"
                            "Please end your continuation with one of these indicators on a new line:\n"
                            "[RESPONSE_COMPLETE] - if your answer is now fully complete\n"
                            "[RESPONSE_INCOMPLETE] - if you still need to continue"
                            """
                            CRITICAL : Don't start the answer with words like "Sure I'll continue the response,
                            just start continuing the answer without any extra introductory paragraphs.
                            """
                        )
                       
                        if rag_prompt:
                            continuation = st.write_stream(generate_response(conv_chain, continuation_prompt, chat_model, True, tab_number))
                        else:
                            continuation = st.write_stream(generate_response(conv_chain, continuation_prompt, chat_model, False, tab_number))
                       
                        response += "\n\n" + continuation
                
                # Clean up any remaining indicators
                response = response.replace("[RESPONSE_COMPLETE]", "").replace("[RESPONSE_INCOMPLETE]", "").strip()
                
                #########################End of Auto Looping Logic##############################################

                st.session_state[history_key].append(f"AI: {response}")

    # Auto-save chat history after each interaction
    save_current_chat_history()


def main():
    if "widget_key" not in st.session_state:
        st.session_state["widget_key"] = str(random.randint(1, 1000000))

    if "file_uploader_key" not in st.session_state:
        st.session_state["file_uploader_key"] = 0

    if "rag_chain" not in st.session_state:
        st.session_state["rag_chain"] = None

    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []

    if "chat_history_gherkin" not in st.session_state:
        st.session_state["chat_history_gherkin"] = []

    if "chat_history_playwright" not in st.session_state:
        st.session_state["chat_history_playwright"] = []

    if "chat_history_code_generator" not in st.session_state:
        st.session_state["chat_history_code_generator"] = []

    if "chat_history_unit_testing" not in st.session_state:
        st.session_state["chat_history_unit_testing"] = []

    if 'question_state' not in st.session_state:
        st.session_state.question_state = False

    if 'fbk' not in st.session_state:
        st.session_state.fbk = str(uuid.uuid4())

    if "current_tab" not in st.session_state:
        st.session_state["current_tab"] = "JSON Payload Test Cases"

    if "model_name" not in st.session_state:
        st.session_state["model_name"] = list(config["models"].keys())[0] if config["models"] else "default"
    
    # Initialize session tracking
    if "session_id" not in st.session_state:
        st.session_state["session_id"] = None
    
    if "chat_history_title" not in st.session_state:
        st.session_state["chat_history_title"] = None
        
    if "chat_history_saved" not in st.session_state:
        st.session_state["chat_history_saved"] = False
    
    result = render_sidebar()

    # Check the length of the result and unpack accordingly
    if len(result) == 2:
        model_kwargs, web_local = result
        memory_window = None
    else:
        model_kwargs, memory_window, web_local = result

    # Store model_kwargs in session state for use in other functions
    st.session_state["model_kwargs"] = model_kwargs

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["JSON Payload Test Cases", "Test Cases (Gherkin Format)", "Playwright Scripts", "Code Generator", "Unit Testing"])

    prompt = st.chat_input()
    
    with tab1:
        uploaded_files = st.file_uploader(
            "Choose a file",
            type=["csv", "txt", "pdf", "py", "docx", "xlsx", "json"],
            accept_multiple_files=True,
            key=st.session_state["file_uploader_key"],
        )

        rag_prompt = ""

        if uploaded_files:
            document_content = ""
            
            for uploaded_file in uploaded_files:
                # Add header for each file
                document_content += f"\n\n=== File: {uploaded_file.name} ===\n"
                
                if uploaded_file.name.endswith('.txt'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content

                elif uploaded_file.name.endswith('.csv'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content

                elif uploaded_file.name.endswith('.pdf'):
                    pdf_reader = PdfReader(uploaded_file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        document_content += page.extract_text()

                elif uploaded_file.name.endswith('.py'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content
                
                elif uploaded_file.name.endswith('.docx'):
                    doc = docx.Document(uploaded_file)
                    file_content = "\n".join([para.text for para in doc.paragraphs])
                    document_content += file_content
                
                elif uploaded_file.name.endswith('.xlsx'):
                    file_content = excel_sheets_to_csv(uploaded_file)
                    document_content += file_content

                elif uploaded_file.name.endswith('.json'):
                    file_content = uploaded_file.read().decode("utf-8")
                    try:
                        json_data = json.loads(file_content)
                        formatted_json = json.dumps(json_data, indent=4)
                        document_content += formatted_json
                    except json.JSONDecodeError:
                        document_content += "Error: Invalid JSON file."

            rag_prompt = document_content
            st.session_state["rag_prompt"] = rag_prompt

        # Display chat history
        if "chat_history" in st.session_state:
            for message in st.session_state["chat_history"]:
                if message.startswith("AI:"):
                    st.chat_message("ai").write(message[3:])
                elif message.startswith("Human:"):
                    st.chat_message("human").write(message[7:])

        if st.session_state["current_tab"] == "JSON Payload Test Cases" and prompt:
            handle_chat_interaction(prompt, "JSON Payload Test Cases", "chat_history", getSystemPrompt, 1)

        if st.session_state.question_state:
            streamlit_feedback(
                feedback_type="thumbs",
                optional_text_label="[Optional]",
                align="flex-start",
                key=st.session_state.fbk+"_tab1_feedback",
                on_submit=fbcb,
            )
            st.session_state.question_state = False


    with tab2:
        uploaded_files = st.file_uploader(
            "Choose a file",
            type=["csv", "txt", "pdf", "py", "docx", "xlsx", "json"],
            accept_multiple_files=True,
            key=st.session_state["file_uploader_key"]+1,
        )

        rag_prompt = ""

        if uploaded_files:
            document_content = ""
            
            for uploaded_file in uploaded_files:
                # Add header for each file
                document_content += f"\n\n=== File: {uploaded_file.name} ===\n"
                
                if uploaded_file.name.endswith('.txt'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content

                elif uploaded_file.name.endswith('.csv'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content

                elif uploaded_file.name.endswith('.pdf'):
                    pdf_reader = PdfReader(uploaded_file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        document_content += page.extract_text()

                elif uploaded_file.name.endswith('.py'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content
                
                elif uploaded_file.name.endswith('.docx'):
                    doc = docx.Document(uploaded_file)
                    file_content = "\n".join([para.text for para in doc.paragraphs])
                    document_content += file_content
                
                elif uploaded_file.name.endswith('.xlsx'):
                    file_content = excel_sheets_to_csv(uploaded_file)
                    document_content += file_content

                elif uploaded_file.name.endswith('.json'):
                    file_content = uploaded_file.read().decode("utf-8")
                    try:
                        json_data = json.loads(file_content)
                        formatted_json = json.dumps(json_data, indent=4)
                        document_content += formatted_json
                    except json.JSONDecodeError:
                        document_content += "Error: Invalid JSON file."

            rag_prompt = document_content
            st.session_state["rag_prompt"] = rag_prompt

        # Display chat history
        if "chat_history_gherkin" in st.session_state:
            for message in st.session_state["chat_history_gherkin"]:
                if message.startswith("AI:"):
                    st.chat_message("ai").write(message[3:])
                elif message.startswith("Human:"):
                    st.chat_message("human").write(message[7:])

        if st.session_state["current_tab"] == "Test Cases (Gherkin Format)" and prompt:
            handle_chat_interaction(prompt, "Test Cases (Gherkin Format)", "chat_history_gherkin", getSystemPromptGherkinFormat, 2)

        if st.session_state.question_state:
            streamlit_feedback(
                feedback_type="thumbs",
                optional_text_label="[Optional]",
                align="flex-start",
                key=st.session_state.fbk+"_gherkin_feedback",
                on_submit=fbcb,
            )
            st.session_state.question_state = False


    with tab3:
        # For Playwright, use chat_history from tab1 as RAG prompt
        rag_prompt_playwright = st.session_state["chat_history"]
        st.session_state["rag_prompt"] = rag_prompt_playwright

        # Display chat history
        if "chat_history_playwright" in st.session_state:
            for message in st.session_state["chat_history_playwright"]:
                if message.startswith("AI:"):
                    st.chat_message("ai").write(message[3:])
                elif message.startswith("Human:"):
                    st.chat_message("human").write(message[7:])

        if st.session_state["current_tab"] == "Playwright Scripts" and prompt:
            handle_chat_interaction(prompt, "Playwright Scripts", "chat_history_playwright", getSystemPromptPlaywright, 3)

        if st.session_state.question_state:
            streamlit_feedback(
                feedback_type="thumbs",
                optional_text_label="[Optional]",
                align="flex-start",
                key=st.session_state.fbk+"_playwright_scripts_feedback",
                on_submit=fbcb,
            )
            st.session_state.question_state = False

    
    with tab4:
        uploaded_files = st.file_uploader(
            "Choose a file",
            type=["csv", "txt", "pdf", "py", "docx", "xlsx", "json"],
            accept_multiple_files=True,
            key=st.session_state["file_uploader_key"]+2,
        )

        rag_prompt = ""

        if uploaded_files:
            document_content = ""
            
            for uploaded_file in uploaded_files:
                # Add header for each file
                document_content += f"\n\n=== File: {uploaded_file.name} ===\n"
                
                if uploaded_file.name.endswith('.txt'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content

                elif uploaded_file.name.endswith('.csv'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content

                elif uploaded_file.name.endswith('.pdf'):
                    pdf_reader = PdfReader(uploaded_file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        document_content += page.extract_text()

                elif uploaded_file.name.endswith('.py'):
                    file_content = uploaded_file.read().decode("utf-8")
                    document_content += file_content
                
                elif uploaded_file.name.endswith('.docx'):
                    doc = docx.Document(uploaded_file)
                    file_content = "\n".join([para.text for para in doc.paragraphs])
                    document_content += file_content
                
                elif uploaded_file.name.endswith('.xlsx'):
                    file_content = excel_sheets_to_csv(uploaded_file)
                    document_content += file_content

                elif uploaded_file.name.endswith('.json'):
                    file_content = uploaded_file.read().decode("utf-8")
                    try:
                        json_data = json.loads(file_content)
                        formatted_json = json.dumps(json_data, indent=4)
                        document_content += formatted_json
                    except json.JSONDecodeError:
                        document_content += "Error: Invalid JSON file."

            rag_prompt = document_content
            st.session_state["rag_prompt"] = rag_prompt

        # Display chat history
        if "chat_history_code_generator" in st.session_state:
            for message in st.session_state["chat_history_code_generator"]:
                if message.startswith("AI:"):
                    st.chat_message("ai").write(message[3:])
                elif message.startswith("Human:"):
                    st.chat_message("human").write(message[7:])

        if st.session_state["current_tab"] == "Code Generator" and prompt:
            handle_chat_interaction(prompt, "Code Generator", "chat_history_code_generator", getSystemPromptCodeGenerator, 4)

        if st.session_state.question_state:
            streamlit_feedback(
                feedback_type="thumbs",
                optional_text_label="[Optional]",
                align="flex-start",
                key=st.session_state.fbk+"_code_generator_feedback",
                on_submit=fbcb,
            )
            st.session_state.question_state = False

    
    with tab5:
        # For Unit Testing, use chat_history_code_generator as RAG prompt
        rag_prompt_unit_testing = st.session_state["chat_history_code_generator"]
        st.session_state["rag_prompt"] = rag_prompt_unit_testing

        # Display chat history
        if "chat_history_unit_testing" in st.session_state:
            for message in st.session_state["chat_history_unit_testing"]:
                if message.startswith("AI:"):
                    st.chat_message("ai").write(message[3:])
                elif message.startswith("Human:"):
                    st.chat_message("human").write(message[7:])

        if st.session_state["current_tab"] == "Unit Testing" and prompt:
            handle_chat_interaction(prompt, "Unit Testing", "chat_history_unit_testing", getSystemPromptUnitTesting, 5)

        if st.session_state.question_state:
            streamlit_feedback(
                feedback_type="thumbs",
                optional_text_label="[Optional]",
                align="flex-start",
                key=st.session_state.fbk+"_unit_testing_feedback",
                on_submit=fbcb,
            )
            st.session_state.question_state = False

    # Auto-save chat history periodically
    if prompt:
        save_current_chat_history()

    print(f"Current Tab = {st.session_state['current_tab']}")
    print(f"Session ID = {st.session_state.get('session_id', 'None')}")
    print(f"Chat History Saved = {st.session_state.get('chat_history_saved', False)}")


if __name__ == '__main__':
    main()